from typing import List, Dict, Tuple, Union, Any
import re
import torch
import fitz
import io
from bs4 import BeautifulSoup
from ..models import Document
import docx
import uuid
from transformers import AutoTokenizer
from sentence_transformers import SentenceTransformer, util
from difflib import SequenceMatcher

class UniversalExtractor():
    def __init__(self):
        pass

    def execute(self, file_bytes: bytes, extension: str) -> List[str]:
        try:
            ext = extension.lower()
            if ext == "pdf":
                return self.extract_from_pdf(file_bytes)
            elif ext == "txt":
                return self.extract_from_txt(file_bytes)
            elif ext == "docx":
                return self.extract_from_docx(file_bytes)
            elif ext in ["html", "htm"]:
                return self.extract_from_html(file_bytes)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
        except Exception as e:
            import logging
            logging.exception(f"preprocessor.execute: Exception {str(e)}")
            raise

    def extract_from_pdf(self, file_bytes: bytes) -> List[str]:
        try:
            buffer = io.BytesIO(file_bytes)
            doc = fitz.open(stream=buffer, filetype="pdf")
            print(doc)
            extracted_blocks = []
            for page in doc:
                blocks = page.get_text("blocks")
                sorted_blocks = sorted(blocks, key=lambda b: (b[1], b[0]))
                for block in sorted_blocks:
                    block_text = block[4].strip()
                    if block_text:
                        cleaned = self.clean_block(block_text)
                        if cleaned:
                            extracted_blocks.append(cleaned)
            return extracted_blocks
        except Exception as e:
            import logging
            logging.exception(f"preprocessor.extract_from_pdf: Exception {str(e)}")
            raise

    def extract_from_txt(self, file_bytes: bytes) -> List[str]:
        try:
            content = file_bytes.decode('utf-8', errors='ignore')
            return content.split("\n\n")
        except Exception as e:
            import logging
            logging.exception(f"preprocessor.extract_from_txt: Exception {str(e)}")
            raise

    def extract_from_docx(self, file_bytes: bytes) -> List[str]:
        try:
            buffer = io.BytesIO(file_bytes)
            doc = docx.Document(buffer)
            extracted_paragraphs = []
            for para in doc.paragraphs:
                block = para.text.strip()
                if block:
                    cleaned = self.clean_block(block)
                    if cleaned:
                        extracted_paragraphs.append(cleaned)
            return extracted_paragraphs
        except Exception as e:
            import logging
            logging.exception(f"preprocessor.extract_from_docx: Exception {str(e)}")
            raise

    def extract_from_html(self, file_bytes: bytes) -> List[str]:
        try:
            raw_html = file_bytes.decode('utf-8', errors='ignore')
            raw_html = raw_html.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
            soup = BeautifulSoup(raw_html, 'html.parser')
            content_tags = soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'li', 'blockquote'])
            extracted_elements = []
            for tag in content_tags:
                block = tag.get_text(separator='\n', strip=True)
                sub_blocks = re.split(r'\n+|•|- ', block)
                for sub in sub_blocks:
                    cleaned = self.clean_block(sub)
                    if cleaned:
                        extracted_elements.append(cleaned)
            return extracted_elements
        except Exception as e:
            import logging
            logging.exception(f"preprocessor.extract_from_html: Exception {str(e)}")
            raise

    def clean_block(self, block: str) -> str:
        block = re.sub(r'\s+', ' ', block)  # collapse excessive whitespace
        block = block.strip()
        return block if len(block) > 5 else ''
    
    def _clean_text(s: str) -> str:
        s = s.replace('\u00A0', ' ')  # non-breaking spaces
        s = re.sub(r'\s+', ' ', s).strip()
        return s


class NoiseRemover:
    """
    Smart Noise Remover and Pre-Cleaner
    Inherits from BasePreprocessor (which inherits from BaseComponent).

    Purpose:
    - Remove headers, footers, page numbers, boilerplate block, and repeated noise
    - Preserve meaningful content, structure, and metadata
    """

    def execute(self, blocks: List[str], min_words: int = 4) -> List[str]:
        print('NoiseRemover called')
        cleaned_blocks = []
        seen = set()

        for block in blocks:
            original = block

            if not block:
                continue
            block = block.strip()

            # Preserve dates like 2025-07-18
            block = re.sub(r'(?<=\d)\s*[-/]\s*(?=\d)', '-', block)

            # Preserve time like 12:04:52 or 12:04
            block = re.sub(r'(?<=\d)\s*[:]\s*(?=\d)', ':', block)

            # Remove common URLs, emails, phone numbers
            block = re.sub(r'\b(?:https?://|www\.)\S+\b', '', block)
            block = re.sub(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '', block)
            # Remove Indian phone numbers
            block = re.sub(r'\b(?:\+91[-\s]?|0)?[6-9]\d{9}\b', '', block)
            # Remove general international-style phone numbers
            block = re.sub(r'\b(?:\+[\d]{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', '', block)

            # Remove leading/trailing HTML/JSON/Tags/brackets etc.
            block = re.sub(r'^<[^>]+>|<[^>]+>$', '', block)
            block = re.sub(r'^[\[{(]+|[\]})]+$', '', block)

            # Clean repeated punctuations but keep date/time/log symbols
            block = re.sub(r'[^\w\s:/\-,]', '', block)  # keeps : / - , for dates and logs
            block = re.sub(r'[^\w\s.,!?()\-+]', '', block) # other noisy symbols

            # Remove extensionra spaces
            block = re.sub(r'\s+', ' ', block).strip()

            # Remove if only digits or random characters
            if re.fullmatch(r'[\d\s]+', block) or re.fullmatch(r'[^\w\s]+', block):
                continue

            # Remove short blocks
            if len(block.split()) < min_words:
                continue

            # Remove near-duplicates (approx match)
            is_duplicate = False
            for seen_block in seen:
                similarity = SequenceMatcher(None, block, seen_block).ratio()
                if similarity > 0.90:
                    is_duplicate = True
                    break

            if is_duplicate:
                continue

            cleaned_blocks.append(block)
            seen.add(block)

        return cleaned_blocks


class SmartAdaptiveChunker():
    def __init__(
        self,
        model_name: str = "bert-base-uncased",
        similarity_threshold: float = 0.80
    ):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.similarity_threshold = similarity_threshold
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')

    def execute(self, blocks: List[str], max_token_len, min_token_len) -> List[Dict[str, Any]]:
        logical_blocks = self.delimiter_split(blocks, max_token_len, min_token_len)
        refined_chunks = self.semantic_refine(logical_blocks, max_token_len)
        return refined_chunks
    

    def delimiter_split(self, blocks: List[str], max_token_len: int, min_token_len: int) -> List[str]:
        print('delimiter_split called')
        chunks = []
        buffer = []
        buffer_tokens = 0

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            block_tokens = len(block.split())

            # If adding this block won't exceed max_token_len, add it to buffer
            if buffer_tokens + block_tokens <= max_token_len:
                buffer.append(block)
                buffer_tokens += block_tokens
            else:
                # Process the buffer first
                if buffer:
                    combined = " ".join(buffer)
                    chunks.extend(self._split_combined_block(combined, max_token_len))
                # Start new buffer
                buffer = [block]
                buffer_tokens = block_tokens

        # Process any remaining buffer
        if buffer:
            combined = " ".join(buffer)
            chunks.extend(self._split_combined_block(combined, max_token_len))

        return chunks

    def _split_combined_block(self, combined: str, max_token_len: int) -> List[str]:
        # Smart sentence splitter (ignores decimals/abbreviations)
        sentences = re.split(r'(?<=[.!?])\s+', combined.strip())

        local_chunks = []
        current_chunk = []
        current_tokens = 0

        for sentence in sentences:
            sentence_tokens = len(sentence.split())

            if sentence_tokens > max_token_len:
                # Flush current chunk if it exists
                if current_chunk:
                    local_chunks.append(' '.join(current_chunk))
                    current_chunk = []
                    current_tokens = 0
                # Add long sentence as its own chunk
                local_chunks.append(sentence.strip())
                continue

            if current_tokens + sentence_tokens <= max_token_len:
                current_chunk.append(sentence.strip())
                current_tokens += sentence_tokens
            else:
                local_chunks.append(' '.join(current_chunk))
                current_chunk = [sentence.strip()]
                current_tokens = sentence_tokens

        if current_chunk:
            local_chunks.append(' '.join(current_chunk))

        return local_chunks


    def make_chunk(self, text: str) -> Document:
        return Document(
            text= text.strip(),
            metadata= {"chunk_id": str(uuid.uuid4()),
                        "token_len": len(self.tokenizer.tokenize(text))}
        )

    def semantic_refine(self, chunks: List[str], max_token_len) -> Tuple[List[Document], Union[torch.Tensor, List[torch.Tensor]]]:
        print("semantic_refine called")

        if not chunks:
            return [], []

        # Encode chunks to get embeddings
        embeddings = self.embedder.encode(chunks, normalize_embeddings=True, convert_to_tensor=True)
        refined_chunks = []
        i = 0
        total_chunks = len(chunks)

        while i < total_chunks:
            current_chunk = chunks[i]
            current_embed = embeddings[i]

            # Try to merge with next chunks if semantically similar
            merged = False
            j = i + 1
            while j < total_chunks:
                sim = util.cos_sim(current_embed, embeddings[j]).item()
                if sim > self.similarity_threshold and len(current_chunk.split() + chunks[j].split()) <= max_token_len:
                    # Merge chunks
                    current_chunk += " " + chunks[j]
                    # Recompute embedding for the merged chunk
                    current_embed = self.embedder.encode(current_chunk, normalize_embeddings=True, convert_to_tensor=True)
                    j += 1
                    merged = True
                else:
                    break

            # Add the final merged chunk
            refined_chunks.append(self.make_chunk(current_chunk))
            i = j if merged else i + 1

        return refined_chunks, embeddings
    
    
    

"""
ok, regarding the context merging issue, I've thought of a method. You know in transformers, we use a positional encoding with the input that tells which neighbour token lies at how much proximity of the current token and the current context length includes before and after tokens of the current token.
So, My plan is to use a proximity based positional encoding with every doc, lets say of length m. So while calculating  the context-similarity of semantically un-matched docs,  we will use positional encoding, along with semantic/cosine similarity, to determine which docs (any length) are very much in contextual similarity with the current doc. Because Im relying on the fact that the words/statements more close to the current statement, will most likely to be related contextually, and more far ones are most likely to relating with other context. and even if a contextual very similar statement is very far in the document, the semantic-similarity score will most  likely score it similar to the current doc.
Correct me if im wrong.

"""


"""

Option 1 — Self-Evaluating / Feedback-Loop RAG Idea: Build an automatic feedback loop where the model evaluates its own answers using an LLM-based “critic” (like a mini DeepSeek-R1 style system).
Option 2 — Adaptive Context Compression Idea: Before sending context to LLM, summarize or compress retrieved docs using another LLM, keeping only the most relevant chunks.
Option 3 — Active Learning from User Feedback Idea: Let users rate answers (thumbs up/down). The system retrains embeddings or adjusts context weighting accordingly.
Option 4 --- Semantic similarity clustering + context merging + redundancy reduction
Option 5 --- Contextual Reasoning Graph (Cognitive Map for RAG) Idea: Convert retrieved docs into nodes and edges, where edges represent reasoning relations like “supports”, “contradicts”, “explains”, “expands”.


"""
